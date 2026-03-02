import os
import re
from datetime import datetime
from typing import Dict, List, Optional
from ..utils.formatting import format_speaker as default_format_speaker_func, format_time

# Check for python-docx availability
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _add_formatted_text(paragraph, text: str):
    """Helper function to add text with markdown bold formatting to a paragraph."""
    # Split by bold markers (**text**)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def export_transcript_to_docx(
    segments: List[Dict],
    output_path: str,
    audio_file: str = None,
    audio_length: float = None,
    format_speaker_func = None
) -> str:
    """
    Export raw transcript from WhisperX to DOCX file.
    """
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Run: pip install python-docx"
    
    format_speaker = format_speaker_func or default_format_speaker_func
    
    doc = Document()
    
    # Title
    title = doc.add_heading('บันทึกการประชุม (Raw Transcript)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_heading('ข้อมูลทั่วไป', level=1)
    
    if audio_file:
        doc.add_paragraph(f"ไฟล์เสียง: {os.path.basename(audio_file)}")
    
    doc.add_paragraph(f"วันที่สร้าง: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    if audio_length:
        mins = int(audio_length // 60)
        secs = int(audio_length % 60)
        doc.add_paragraph(f"ความยาวเสียง: {mins}:{secs:02d} นาที")
    
    doc.add_paragraph()
    
    # Transcript content
    doc.add_heading('เนื้อหาการประชุม', level=1)
    
    # Add table for transcript
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'เวลาเริ่ม'
    header_cells[1].text = 'เวลาจบ'
    header_cells[2].text = 'ผู้พูด'
    header_cells[3].text = 'ข้อความ'
    
    # Make headers bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add segments
    for segment in sorted(segments, key=lambda x: x.get('start', 0)):
        row = table.add_row().cells
        
        start = segment.get('start', 0)
        end = segment.get('end', 0)
        # Use speaker name as-is (already mapped by pipeline)
        speaker = segment.get('speaker', '') or ''
        if not speaker or speaker.startswith('SPEAKER_'):
            speaker = format_speaker(segment.get('speaker'))
        text = segment.get('text', '').strip()
        
        row[0].text = format_time(start)
        row[1].text = format_time(end)
        row[2].text = speaker
        row[3].text = text
    
    # Add Combined Text section
    doc.add_paragraph()
    doc.add_heading('เนื้อหารวม (Combined Text)', level=1)
    
    # Build combined text with speaker labels
    combined_lines = []
    current_speaker = None
    current_text = []
    
    for segment in sorted(segments, key=lambda x: x.get('start', 0)):
        # Use speaker name as-is (already mapped by pipeline)
        speaker = segment.get('speaker', '') or ''
        if not speaker or speaker.startswith('SPEAKER_'):
            speaker = format_speaker(segment.get('speaker'))
        text = segment.get('text', '').strip()
        
        if speaker == current_speaker:
            current_text.append(text)
        else:
            if current_speaker and current_text:
                combined_lines.append(f"[{current_speaker}]: {' '.join(current_text)}")
            current_speaker = speaker
            current_text = [text]
    
    # Add last speaker's text
    if current_speaker and current_text:
        combined_lines.append(f"[{current_speaker}]: {' '.join(current_text)}")
    
    # Add combined text to document
    combined_text = "\n\n".join(combined_lines)
    p = doc.add_paragraph(combined_text)
    
    # Save document
    doc.save(output_path)
    return output_path


def export_summary_to_docx(
    summary_text: str,
    output_path: str,
    speaker_summary: Dict = None,
    meeting_type_id: int = 0
) -> str:
    """
    Export AI summary to a formatted DOCX file with participant header section.
    
    Args:
        summary_text: The AI-generated summary text
        output_path: Path to save the DOCX file
        speaker_summary: Dictionary containing 'speaking_time' and 'word_count' per speaker
        meeting_type_id: Meeting type ID for position formatting
    """
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Run: pip install python-docx"
    
    # Position templates based on meeting type
    POSITION_TEMPLATES = {
        0: {"leader": "ประธาน", "main": "ผู้นำเสนอ", "participant": "ผู้เข้าร่วม"},
        1: {"leader": "ประธาน", "main": "กรรมการ", "participant": "ผู้ถือหุ้น"},  # Shareholder
        2: {"leader": "ประธาน", "main": "กรรมการ", "participant": "ผู้เข้าร่วม"},  # Board
        3: {"leader": "ผู้จัดการโครงการ", "main": "ผู้รับผิดชอบ", "participant": "ทีมงาน"},  # Planning
        4: {"leader": "ผู้รายงาน", "main": "ผู้รับผิดชอบ", "participant": "ผู้เข้าร่วม"},  # Progress
        5: {"leader": "ผู้บริหาร", "main": "ผู้นำเสนอ", "participant": "ผู้เข้าร่วม"},  # Strategy
        6: {"leader": "หัวหน้าทีม", "main": "ผู้เกี่ยวข้อง", "participant": "ผู้เข้าร่วม"},  # Incident
        7: {"leader": "ผู้แทนบริษัท", "main": "ผู้นำเสนอ", "participant": "ลูกค้า"},  # Client
        8: {"leader": "ผู้บรรยาย", "main": "ผู้ช่วยบรรยาย", "participant": "ผู้เข้าร่วม"},  # Workshop
        9: {"leader": "ประธาน", "main": "ผู้บริหาร", "participant": "ผู้เข้าร่วม"},  # Executive
        10: {"leader": "หัวหน้าทีม", "main": "ผู้นำเสนอ", "participant": "สมาชิกทีม"},  # Team
        11: {"leader": "ประธาน", "main": "ผู้นำเสนอ", "participant": "ผู้เข้าร่วม"},  # General
    }
    
    doc = Document()
    
    # Title
    title = doc.add_heading('สรุปการประชุม', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # ============ PARTICIPANT HEADER SECTION ============
    if speaker_summary:
        speakers_time = speaker_summary.get('speaking_time', {})
        total_time = sum(speakers_time.values()) if speakers_time else 1
        
        if speakers_time:
            # Get position template for this meeting type
            positions = POSITION_TEMPLATES.get(meeting_type_id, POSITION_TEMPLATES[0])
            
            # Add numbered heading
            heading = doc.add_paragraph()
            heading_run = heading.add_run('1. ผู้เข้าร่วมประชุมและตำแหน่ง')
            heading_run.bold = True
            heading_run.font.size = Pt(14)
            
            # Sort speakers by speaking time (most active first)
            sorted_speakers = sorted(speakers_time.items(), key=lambda x: -x[1])
            
            # Group by position based on speaking percentage
            leader_added = False
            main_speakers = []
            participants = []
            
            for idx, (speaker, time_sec) in enumerate(sorted_speakers):
                pct = (time_sec / total_time * 100) if total_time > 0 else 0
                
                if idx == 0 and pct > 25:  # First speaker with significant time = leader
                    # Add leader with actual speaker name
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{positions['leader']}: ").bold = True
                    p.add_run(speaker)
                    leader_added = True
                elif pct > 10:  # Significant speaker = main role
                    main_speakers.append(speaker)
                else:  # Regular participant
                    participants.append(speaker)
            
            # If no leader was added, use first speaker
            if not leader_added and sorted_speakers:
                first_speaker = sorted_speakers[0][0]
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{positions['leader']}: ").bold = True
                p.add_run(first_speaker)
            
            # Add main speakers (กรรมการ/ผู้นำเสนอ)
            if main_speakers or participants:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{positions['main']}:").bold = True
                
                all_others = main_speakers + participants
                for speaker in all_others:
                    sub_p = doc.add_paragraph(style='List Bullet 2')
                    sub_p.add_run(speaker)
            
            doc.add_paragraph()  # Spacer after participant section
    
    # ============ END PARTICIPANT HEADER SECTION ============
    
    # Header mapping for structured sections
    section_headers = {
        'ประเภท': 1,
        'ผู้เข้าร่วมประชุม': 1,
        'สรุปการประชุม': 1,
        'การสั่งงาน': 1,
        'มอบหมาย': 1,
        'คำถามสำคัญ': 1,
        'ข้อตกลง': 1,
        'มติ': 1,
        'สถานะ': 2,
        'ความคืบหน้า': 2,
        'ปัญหา': 2,
        'แนวทางแก้': 2,
        'งานถัดไป': 2,
        'นโยบาย': 2,
        'การอนุมัติ': 2,
    }
    
    def remove_emoji(text):
        """Remove emoji from text"""
        emoji_pattern = re.compile("["
            u"\U0001F600-\U0001F64F"  # emoticons
            u"\U0001F300-\U0001F5FF"  # symbols & pictographs
            u"\U0001F680-\U0001F6FF"  # transport & map symbols
            u"\U0001F1E0-\U0001F1FF"  # flags
            u"\U00002702-\U000027B0"
            u"\U000024C2-\U0001F251"
            u"\U0001f926-\U0001f937"
            u"\U00010000-\U0010ffff"
            "]+", flags=re.UNICODE)
        return emoji_pattern.sub('', text).strip()
    
    def get_header_level(text):
        """Determine header level based on content"""
        clean_text = remove_emoji(text)
        for keyword, level in section_headers.items():
            if keyword in clean_text:
                return level
        return None
    
    # Parse markdown and add to document
    lines = summary_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Remove emoji from line
        clean_line = remove_emoji(line)
        
        # Handle headers (**, ##, etc.)
        if line.startswith('**') and line.endswith('**'):
            # Bold header - check if it's a section header
            text = line.strip('*').strip()
            clean_text = remove_emoji(text)
            header_level = get_header_level(text)
            
            if header_level:
                doc.add_heading(clean_text, level=header_level)
            else:
                p = doc.add_paragraph()
                run = p.add_run(clean_text)
                run.bold = True
                run.font.size = Pt(12)
        elif line.startswith('##'):
            text = remove_emoji(line.lstrip('#').strip())
            doc.add_heading(text, level=2)
        elif line.startswith('#'):
            text = remove_emoji(line.lstrip('#').strip())
            doc.add_heading(text, level=1)
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            text = remove_emoji(line[2:].strip())
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            _add_formatted_text(p, clean_line)
    
    # Save document
    doc.save(output_path)
    return output_path


def export_both(
    segments: List[Dict],
    summary_text: str,
    base_path: str,
    audio_file: str = None,
    audio_length: float = None,
    format_speaker_func = None,
    output_dir: str = "doc",
    speaker_summary: Dict = None,
    meeting_type_id: int = 0
) -> Dict[str, str]:
    """
    Export both transcript and summary to DOCX files.
    
    Args:
        segments: Transcription segments from WhisperX
        summary_text: AI-generated summary text
        base_path: Base path/filename for output files
        audio_file: Original audio file path
        audio_length: Length of audio in seconds
        format_speaker_func: Custom speaker formatting function
        output_dir: Directory to save DOCX files
        speaker_summary: Dictionary with 'speaking_time' and 'word_count' per speaker
        meeting_type_id: Meeting type ID for position formatting
    """
    # Ensure output directory exists
    if not os.path.isabs(output_dir):
        # If relative path, make it relative to current working directory
        output_dir = os.path.join(os.getcwd(), output_dir)
    
    os.makedirs(output_dir, exist_ok=True)
    
    # Get base filename from base_path
    base_filename = os.path.basename(base_path)
    
    transcript_path = os.path.join(output_dir, f"{base_filename}_transcript.docx")
    summary_path = os.path.join(output_dir, f"{base_filename}_summary.docx")
    
    results = {}
    
    # Export transcript
    results['transcript'] = export_transcript_to_docx(
        segments=segments,
        output_path=transcript_path,
        audio_file=audio_file,
        audio_length=audio_length,
        format_speaker_func=format_speaker_func
    )
    
    # Export summary with speaker info and meeting type for participant header
    results['summary'] = export_summary_to_docx(
        summary_text=summary_text,
        output_path=summary_path,
        speaker_summary=speaker_summary,
        meeting_type_id=meeting_type_id
    )
    
    return results
